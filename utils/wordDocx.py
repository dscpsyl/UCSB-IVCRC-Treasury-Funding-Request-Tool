import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

#? Searches through a document for instances and replaces them and then returns the document
def replaceTextPar(doc, searchText, replaceText, file):
    document = Document(doc)
    
    for paragraph in document.paragraphs:
        if searchText in paragraph.text:
            origText = paragraph.text
            newText = origText.replace(searchText, replaceText)
            paragraph.text = newText
            
    document.save("env/" + file + ".docx")
    return 0

#? Searches through a document's tables for instances and replaces them and then returns the document
def replaceTextTable(doc, searchText, replaceText, file):
    document = Document(doc)
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                origText = cell.text
                newText = origText.replace(searchText, replaceText)
                cell.text = newText
                
    document.save("env/" + file + ".docx")
    return 0

#? Adds a hyperlink to the end of a paragraph
def addHyperlink(paragraph, url, text, color="2854C5", underline="single"):

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    # Underline types are specified in docx.enum.text.WD_UNDERLINE
    u = docx.oxml.shared.OxmlElement('w:u')
    if not underline:
      u.set(docx.oxml.shared.qn('w:val'), 'none')
    else:
      u.set(docx.oxml.shared.qn('w:val'), underline)
    rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

#? Adds a hyperlink to the specific part that the Funding Agreement needs
def theMostNicheFunctionBecauseIDontWantMainToLookTooUgly(doc, frId, file):
    document = docx.Document(doc)
    
    for t in document.tables:
        for r in t.rows:
            for c in r.cells:
                oText = c.text
                if oText == "[frReport]":
                    c.text = ""
                    p = c.paragraphs[0]
                    p.add_run("Expected end of grant and ")
                    addHyperlink(p, 'https://docs.google.com/document/d/' + frId, '[Follow-up Report]')
                    p.add_run(" due. See details below.")
    
    document.save("env/" + file + ".docx")
    return 0


#? takes in a search and highlight paragraph containing that search
def highlightPar(doc, searchText, file):
    document = Document(doc)
    
    found = False
    
    for paragraph in document.paragraphs:
        if searchText in paragraph.text:
            found = True
            for r in paragraph.runs:
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
    document.save("env/" + file + ".docx")
    return found